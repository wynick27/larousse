#!/usr/bin/env python3
"""
NiceGUI-based web proofing tool with configurable fields and data sources.
Configuration is loaded from ./data/config.json by default, or switched via toolbar.
Supports showing multiple records at once (configurable).
"""

from __future__ import annotations

import asyncio
import json
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Tuple
import re
from abc import ABC, abstractmethod
from nicegui import ui, app
from fastapi.responses import Response
from difflib import SequenceMatcher
from PIL import Image, ImageOps
import fitz
import io

# =============================
# Global Config State
# =============================
CONFIG_PATH = Path('./data/config.json')
CONFIG: dict = {}

OUTPUT_PATH: Path = Path('')
IMAGE_CONFIG: dict = None
CANDIDATES_CONFIG: Dict = {}
KEYNAME: str = 'no'
FIELDS: List[str] = []
FIELD_MAP: Dict[str, dict] = {}
FIELD_LABELS: Dict[str, str] = {}
FIELD_MODES: Dict[str, str] = {}
FIELD_PDF_KEYS: Dict[str, str] = {}
TITLE: str = 'JSON 文本校对工具 (Larousse)'
ITEMS_PER_PAGE: int = 1
AUTOSAVE_SECONDS = 60
IMAGENAME = None

# 配置 row_mode
ROW_MODE = bool(CONFIG.get('row_mode', False))
pdf_doc = None
candidates: dict = {}

store: BaseStore
current_filter = None
filtered_indices = []
main_column = None
record_status_label = None
compact_container = None
timer = None
image_source: ImageSource = None

def get_config_files():
    config_dir = Path('./config')
    return [str(p) for p in config_dir.glob('*.json')]

# =============================
# Config Loader
# =============================

def load_config(path: Path):
    global CONFIG, OUTPUT_PATH, IMAGE_CONFIG, CANDIDATES_CONFIG, KEYNAME,ROW_MODE,AUTOSAVE_SECONDS
    global FIELDS, FIELD_LABELS, FIELD_MODES, FIELD_PDF_KEYS, pdf_doc, candidates, IMAGENAME, ITEMS_PER_PAGE, store, timer
    global FIELD_MAP
    with open(path, 'r', encoding='utf-8') as f:
        CONFIG = json.load(f)

    IMAGE_CONFIG = CONFIG.get('images')
    CANDIDATES_CONFIG = CONFIG['candidates']
    KEYNAME = CONFIG['key_name']
    ITEMS_PER_PAGE = int(CONFIG.get('items_per_page', 1))

    FIELD_MAP = CONFIG['field_config']
    FIELDS = list(FIELD_MAP.keys())
    FIELD_LABELS = {k: v['label'] for k, v in FIELD_MAP.items()}
    FIELD_MODES = {k: v.get('mode', 'normal') for k, v in FIELD_MAP.items()}
    FIELD_PDF_KEYS = {k: v.get('pdf_key') for k, v in FIELD_MAP.items()}

    # load PDF
    #if PDF_PATH and Path(PDF_PATH).exists():
    #    pdf_doc = fitz.open(PDF_PATH)
    #else:
    #    pdf_doc = None

    # load candidates

    if isinstance(CANDIDATES_CONFIG, dict):
        candidates = {}
        for key, path in CANDIDATES_CONFIG.items():
            candidate_path = Path(path)
            if candidate_path.exists():
                with open(candidate_path, 'r', encoding='utf-8') as f:
                    candidates[key] = json.load(f)

    AUTOSAVE_SECONDS = int(CONFIG.get('autosave_seconds', 60))
    if timer is not None:
        timer.delete()
    if AUTOSAVE_SECONDS > 0:
        timer = ui.timer(AUTOSAVE_SECONDS, save_now)
    ROW_MODE = bool(CONFIG.get('row_mode', True))

    input_file = CONFIG.get('input_file')
    output_file = CONFIG.get('output_file', input_file)
    store = create_store(input_file, KEYNAME)
    store.load()
# =============================
# Data Layer
# =============================

def safe_read_json(path):
    if isinstance(path, str):
        path = Path(path)
    with path.open('r', encoding='utf-8') as f:
        return json.load(f)


def safe_write_json(path: Path, data: List[dict]) -> None:
    tmp = path.with_suffix('.tmp')
    with tmp.open('w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    tmp.replace(path)

class ImageSource:
    def __init__(self, cfg: dict):
        self.type = cfg.get('type', 'dir')
        self.path = Path(cfg['path'])
        self.cfg = cfg
        if self.type == 'dir':
            if not self.path.exists() or not self.path.is_dir():
                raise ValueError(f'图片目录不存在: {self.path}')
        elif self.type == 'pdf':
            if not self.path.exists() or not self.path.is_file():
                raise ValueError(f'图片PDF文件不存在: {self.path}')
            self.pdf_doc = fitz.open(self.path)
        else:
            raise ValueError(f'不支持的图片类型: {self.type}')

    def get_image(self, key: str) -> Image.Image | None:
        if self.type == 'dir':
            img_path = self.path / key
            if img_path.exists() and img_path.is_file():
                try:
                    img = Image.open(img_path)
                    img = ImageOps.exif_transpose(img)
                    return img
                except Exception as e:
                    print(f'无法打开图片 {img_path}: {e}')
                    return None
            else:
                return None
        elif self.type == 'pdf':
            try:
                page_num = int(key)
                if 0 <= page_num < len(self.pdf_doc):
                    page = self.pdf_doc[page_num]
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    return img
                else:
                    return None
            except Exception as e:
                print(f'无法从PDF获取图片 {key}: {e}')
                return None
        return None


class BaseStore(ABC):
    def __init__(self, path_or_cfg, key_name):
        if isinstance(path_or_cfg, str):
            self.path = Path(path_or_cfg)
            self.cfg = {}
        else:
            self.path = path_or_cfg['path']
            self.cfg = path_or_cfg
        self.key_name = key_name
        self.records = []
        self.dirty = False

    @abstractmethod
    def load(self):
        ...

    @abstractmethod
    def save(self,path = None):
        ...

    def set_field(self, index: int, field_name: str, value: str):
        self.records[index][field_name] = value
        self.dirty = True

class JsonStore(BaseStore):
    def load(self):
        self.records = safe_read_json(self.path)

    def save(self,path = None):
        if self.dirty:
            safe_write_json(self.path, self.records)
            self.dirty = False

class XlsxStore(BaseStore):
    def load(self):
        import openpyxl
        wb = openpyxl.load_workbook(self.path)
        sheet = wb.active
        headers = [c.value for c in sheet[1]]
        self.records = []
        for row in sheet.iter_rows(min_row=2, values_only=True):
            rec = dict(zip(headers, row))
            self.records.append(rec)

    def save(self,path = None):
        import openpyxl
        if not self.dirty:
            return
        wb = openpyxl.Workbook()
        sheet = wb.active
        if not self.records:
            return
        headers = list(self.records[0].keys())
        sheet.append(headers)
        for rec in self.records:
            sheet.append([rec.get(h, "") for h in headers])
        wb.save(self.path)
        self.dirty = False

class DocxStore(BaseStore):
    def load(self):
        import docx
        doc = docx.Document(self.path)
        self.records = []
        for i, para in enumerate(doc.paragraphs):
            self.records.append({self.key_name: i+1, 'text': para.text})

    def save(self,path = None):
        import docx
        if not self.dirty:
            return
        doc = docx.Document()
        for rec in self.records:
            doc.add_paragraph(rec.get('text', ''))
        doc.save(self.path)
        self.dirty = False

class PdfStore(BaseStore):
    def load(self):
        pdf = fitz.open(self.path)
        self.records = []
        for i, page in enumerate(pdf):
            self.records.append({self.key_name: i+1, 'text': page.get_text(), 'page':i+1})

    def save(self,path = None):
        # PDF 保存复杂，默认导出为txt
        if not self.dirty:
            return
        txt_path = self.path.with_suffix('.out.txt')
        with txt_path.open('w', encoding='utf-8') as f:
            for rec in self.records:
                f.write(rec.get('text', '') + '\n\n')
        self.dirty = False

class TxtStore(BaseStore):
    def __init__(self, cfg, key_name):
        super().__init__(cfg['path'], key_name)
        self.start_pat = re.compile(cfg.get('start', r'^'), re.M)
        self.end_pat = re.compile(cfg.get('end', r'\Z'), re.M)
        self.split_pat = re.compile(cfg.get('split', r'\n'), re.M)
        self.page_pat =  re.compile(cfg.get('page'), re.M) if cfg.get('page') else None
        self.type = cfg.get('type', 'simple')
        self.fields_cfg = cfg.get('fields', {})
        self.grammar_path = cfg.get('grammar_path')
        self.parser = None
        if self.type == 'grammar' and self.grammar_path:
            grammar_text = Path(self.grammar_path).read_text(encoding='utf-8')
            self.parser = Lark(grammar_text, parser='lalr')

        self.prefix = ''
        self.suffix = ''

    def load(self):
        text = Path(self.path).read_text(encoding='utf-8')

        # 定位 start / end
        start_match = self.start_pat.search(text)
        end_match = self.end_pat.search(text)
        start_idx = start_match.end() if start_match else 0
        end_idx = end_match.start() if end_match else len(text)

        self.prefix = text[:start_idx]
        self.suffix = text[end_idx:]
        core_text = text[start_idx:end_idx]

        # 分割 + 保留分隔符
        self.records = []
        last_pos = 0
        cur_page = 0
        cur_num = 0

        for i, match in enumerate(self.split_pat.finditer(core_text)):
            segment = core_text[last_pos:match.start()]
            if self.page_pat and self.page_pat.match(segment):
                match = self.page_pat.match(segment)
                if cur_page != int(match.group(1)):
                    cur_num = 0
                cur_page = int(match.group(1))
                continue
            rec = {"no_page": cur_num+1, '_raw': segment, '_separator': match.group(0)}
            self.records.append(rec)
            if self.page_pat:
                rec['page'] = cur_page
            last_pos = match.end()
            cur_num += 1

        # 最后一个片段
        if last_pos < len(core_text):
            rec = {"no_page": cur_num+1, '_raw': core_text[last_pos:], '_separator': ''}
            if self.page_pat:
                rec['page'] = cur_page
            self.records.append(rec)

        for i, rec in enumerate(self.records):
            seg = rec['_raw']
            if self.type == 'grammar' and self.parser:
                
                tree: Tree = self.parser.parse(seg)
                rec['_tree'] = tree
                for child in tree.children:
                    if isinstance(child, Tree):
                        rec[child.data] = seg[child.meta.start_pos:child.meta.end_pos]
            elif self.type == 'simple':
                for fname, regex in self.fields_cfg.items():
                    for m in re.finditer(regex, seg, re.M):
                        rec.setdefault('_matches', {})[fname] = (m.start(), m.end(), m.group(0))
                        rec[fname] = m.group(1) if m.groups() else m.group(0)
                rec['text'] = seg
            self.records.append(rec)

    def set_field(self, index: int, field_name: str, value: str):
        rec = self.records[index]
        if self.type == 'simple' and '_matches' in rec and field_name in rec['_matches']:
            start, end, old = rec['_matches'][field_name]
            raw = rec['_raw']
            rec['_raw'] = raw[:start] + value + raw[end:]
            rec['_matches'][field_name] = (start, start+len(value), value)
        elif self.type == 'grammar' and '_tree' in rec:
            # 找到对应子树位置替换
            tree: Tree = rec['_tree']
            new_text = list(rec['_raw'])
            for child in tree.children:
                if isinstance(child, Tree) and child.data == field_name:
                    s, e = child.meta.start_pos, child.meta.end_pos
                    new_text[s:e] = value
                    child.meta.start_pos, child.meta.end_pos = s, s+len(value)
            rec['_raw'] = ''.join(new_text)
        else:
            rec['_raw'] = value
        rec[field_name] = value
        self.dirty = True

    def save(self,path = None):
        if not self.dirty:
            return
        rebuilt_segments = [rec['_raw'] for rec in self.records]
        out = []
        for i, seg in enumerate(rebuilt_segments):
            out.append(seg)
            if i < len(self.separators):
                out.append(self.separators[i])
        final_text = self.prefix + ''.join(out) + self.suffix
        Path(self.path).write_text(final_text, encoding='utf-8')
        self.dirty = False

class LarousseTxtStore(BaseStore):
    def load(self):
        text = Path(self.path).read_text(encoding='utf-8')

        start_pos = text.find('〈1〉')
        
        self.prefix = text[:start_pos]
        core_text = text[start_pos:]
        self.lines = core_text.splitlines()

        # 分割 + 保留分隔符
        self.records = []

        words = []
        cur_word = None
        cur_page = 0
        page_start = False
        
        cur_no = 0
        cur_page_no = 0
        for line_no,line in enumerate(self.lines):
            if not line.strip():
                continue
            if match := re.match('〈(\d+)〉',line):
                cur_page = int(match.group(1))
                page_start = True
                cur_page_no = 0
                
            elif page_start and re.fullmatch(r'(?i)^\*?[a-zàâçéèêëîïöôûùüÿñæœ \.\-\']+',line):
                pass
            elif re.match(r'(?i)^(\d+\.\s*)?\*?[a-zàâçéèêëîïôöûùüÿñæœ \.,\-\']+(\(.{2,10}\)\s*)?\[|^[a-zàâçéèêëîïôöûùüÿñæœ\-]+,?\s*(préfixe|préf.)|[A-Z][a-zàâçéèêëîïôöûùüÿñæœ]+\s*\([a-zàâçéèêëîïôöûùüÿñæœ ]+\)',line) or\
                not page_start:
                cur_no += 1
                cur_page_no += 1
                word = {'text': line, 'page': cur_page, 'line':line_no, 'no': cur_no, 'id':f"{cur_page}.{cur_page_no}"}
                cur_word = word
                self.records.append(word)
                page_start = False
            elif '→' in line:
                cur_no += 1
                cur_page_no += 1
                word = {'text': line, 'page': cur_page, 'line':line_no, 'no': cur_no, 'id':f"{cur_page}.{cur_page_no}"}
                cur_word = word
                self.records.append(word)
                page_start = False
            else:
                cur_word['text'] += '\n' + line
                if isinstance(cur_word['page'],int):
                    cur_word['page'] = [cur_word['page']]
                    cur_word['line'] = [cur_word['line']]
                if not cur_page in cur_word['page']:
                    cur_word['page'].append(cur_page)
                cur_word['line'].append(line_no)
                page_start = False


    def save(self,path = None):
        if not self.dirty:
            return
        for rec in self.records:
            if isinstance(rec['line'],list):
                for line_no,text in zip(rec['line'],rec['text'].splitlines()):
                    self.lines[line_no] = rec['text']
            else:
                self.lines[rec['line']] = rec['text']
        final_text = self.prefix + '\n'.join(self.lines)
        Path(self.path).write_text(final_text, encoding='utf-8')
        self.dirty = False

store_extension_map = {
    '.json': JsonStore,
    '.xlsx': XlsxStore,
    '.docx': DocxStore,
    '.pdf': PdfStore,
    '.txt': TxtStore,
}

store_type_map = {
    'larousse': LarousseTxtStore,
}

def create_store(path_or_cfg, key_name: str) -> BaseStore:
    if isinstance(path_or_cfg, str):
        path = path_or_cfg
    else:
        path = path_or_cfg['path']
    ext = Path(path).suffix.lower()
    type = path_or_cfg.get('type') if isinstance(path_or_cfg, dict) else None
    if type in store_type_map:
        return store_type_map[type](path_or_cfg, key_name)
    elif ext in store_extension_map:
        return store_extension_map[ext](path_or_cfg, key_name)
    else:
        raise ValueError(f'不支持的文件类型: {ext}')

def check_brackets(s: str):
    # 定义所有括号的对应关系
    brackets = {
        '(': ')', '[': ']', '{': '}',
        '（': '）', '【': '】', '｛': '｝',
        '《': '》', '「': '」', '『': '』',
        '〔': '〕', '〖': '〗',
        '⟪': '⟫'
    }

    opening = set(brackets.keys())
    closing = set(brackets.values())
    stack = []

    for idx, char in enumerate(s, start=1):  # 下标从1开始，便于提示
        if char in opening:  # 左括号入栈
            stack.append((char, idx))
        elif char in closing:  # 遇到右括号
            if not stack:
                return f"第 {idx} 个字符 '{char}' 没有匹配的左括号"
            last, pos = stack.pop()
            if brackets[last] != char:
                return f"第 {idx} 个字符 '{char}' 与第 {pos} 个字符 '{last}' 不匹配"

    if stack:
        last, pos = stack[-1]
        return f"第 {pos} 个字符 '{last}' 没有匹配的右括号"

    return True

def make_filter(f_cfg):
    """返回一个可调用对象 filter_func(record)->bool"""
    ftype = f_cfg.get("type")
    if ftype == "regex":
        field = f_cfg["field"]
        pattern = re.compile(f_cfg["pattern"])
        return lambda r: bool(pattern.search(str(r.get(field, ""))))
    elif ftype == "paren_match":
        field = f_cfg["field"]
        return lambda r: check_brackets(str(r.get(field, ""))) == True
    elif ftype == "lambda":
        # 用 eval 动态构造函数（注意安全，配置要可信）
        code = f_cfg["code"]
        func = eval(code)
        if not callable(func):
            raise ValueError("lambda filter must be callable")
        return func
    elif ftype == "candidate_half_mismatch":
        # 特殊：检查候选来源是否至少一半不一致
        def _f(r):
            from difflib import SequenceMatcher
            result = []
            for field in FIELDS:
                cands = get_candidate(r[KEYNAME], field)
                if cands:
                    base = r.get(field, "")
                    for v in cands.values():
                        sm = SequenceMatcher(None, base, v)
                        if sm.ratio() < 0.5:
                            result.append(False)
                        else:
                            result.append(True)
            # True 表示合格 -> 所以要取返回 False 的
            return all(result) if result else True
        return _f
    else:
        raise ValueError(f"未知过滤器类型: {ftype}")


def apply_filter(filter_func):
    global filtered_indices, current_index
    filtered_indices = [
        i for i, r in enumerate(store.records) if not filter_func(r)
    ]
    current_index = 0
    refresh_record_view()

def clear_filter():
    global filtered_indices, current_index
    filtered_indices = []
    current_index = 0
    refresh_record_view()

def get_visible_index(idx):
    """返回当前 idx 对应的实际 records 索引"""
    if filtered_indices:
        return filtered_indices[idx]
    return idx

def get_visible_count():
    return len(filtered_indices) if filtered_indices else len(store.records)

def update_record_status():
    total = get_visible_count()
    if total == 0:
        record_status_label.text = "无记录"
    elif ITEMS_PER_PAGE == 1:
        record_status_label.text = f"第 {current_index+1} 条 / 共 {total} 条"
    else:
        # current_index 是相对过滤后的索引
        record_status_label.text = f"第 {current_index+1}-{min(current_index+ITEMS_PER_PAGE,total)} 条 / 共 {total} 条"
# =============================
# Candidate Provider
# =============================

def get_candidate(rec_key: any, field_name: str) -> Dict[str, str]:
    result = {}
    for name, candidate in candidates.items():
        entry = candidate.get(rec_key)
        if entry and field_name in entry:
            result[name] = entry[field_name]
    return result

# =============================
# Diff Engine (character-level, no tokenization)
# =============================

def opcodes_for_diff(a: str, b: str,ignore_pattern = None) -> List[Tuple[str, int, int, int, int]]:
    sm = SequenceMatcher(a=a, b=b, autojunk=False)
    ops = sm.get_opcodes()
    filtered_ops = []
    if ignore_pattern is None:
        return ops
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        a_chunk = sm.a[i1:i2]
        b_chunk = sm.b[j1:j2]
        # 如果差异完全由 ignore_re 匹配的字符组成，直接当作 equal
        if tag != "equal" and re.sub(ignore_pattern,"",a_chunk) == re.sub(ignore_pattern,"",b_chunk):
            filtered_ops.append(("equal", i1, i2, j1, j2))
        else:
            filtered_ops.append((tag, i1, i2, j1, j2))
    return filtered_ops


def apply_single_chunk(a: str, b: str, opcode_index: int, action: str = 'auto', ignore_pattern=None) -> str:
    ops = opcodes_for_diff(a,b,ignore_pattern)
    parts: List[str] = []
    for idx, (tag, i1, i2, j1, j2) in enumerate(ops):
        if tag == 'equal':
            parts.append(a[i1:i2])
        elif idx == opcode_index:
            if tag == 'insert':
                if action in ('insert', 'replace', 'auto'):
                    parts.append(b[j1:j2])
            elif tag == 'delete':
                if action in ('delete', 'auto'):
                    pass
                else:
                    parts.append(a[i1:i2])
            elif tag == 'replace':
                if action != 'delete':
                    parts.append(b[j1:j2])
        else:
            if tag in ('replace', 'delete'):
                parts.append(a[i1:i2])
    return ''.join(parts)

# =============================
# UI State & Styles
# =============================
current_index = 0

ui.add_head_html("""
    <style>
      .chunk { display:inline-block; margin: 0 1px; padding: 0 2px; border-radius: 4px; cursor: pointer; user-select: none; white-space: pre-wrap; }
      .chunk.equal { background: transparent; color: inherit; cursor: default; }
      .chunk.insert { background-color: #d4f8d4; }
      .chunk.delete { background-color: #f8d4d4; text-decoration: line-through; }
      .muted { opacity: 0.7; }
      .mono { font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, monospace; }
      .flex-break { flex-basis: 100%; height: 0; }
      .newline-glyph { opacity: .6; }
      .config-btn { background-color: #eee; color: #333; padding: 4px 8px; border-radius: 4px; cursor: pointer; }
    </style>
""")

field_editors: Dict[Tuple[int,str], ui.element] = {}
field_diff_containers: Dict[Tuple[int,str], ui.element] = {}

# =============================
# Save & Autosave
# =============================

def save_now(_: object = None) -> None:
    store.save(OUTPUT_PATH)
    if store.dirty is False:
        ui.notify('已保存 ✔')


app.on_shutdown(save_now)

# =============================
# Navigation & Updates
# =============================

def set_current_index(new_index: int) -> None:
    global current_index
    max_idx = get_visible_count() - 1
    current_index = max(0, min(new_index, max_idx))
    refresh_record_view()


def update_field(index:int, field_name: str, value: str) -> None:
    store.set_field(index, field_name, value)
    if FIELD_MODES[field_name] == 'normal':
        render_diffs_for_field(index, field_name)
    else:
        render_compact_fields(compact_container)

# =============================
# Rendering helpers
# =============================

def _emit_text_with_breaks(text: str, cls: str, on_click=None, show_newline_glyph: bool = False):
    """Render text; whenever a "\n" is encountered, force a new flex line by inserting a flex-break.
    If show_newline_glyph is True, also render a small clickable glyph before breaking the line.
    """
    tokens = re.split(r"(\n)", text)
    for tok in tokens:
        if tok == "":
            continue
        if tok == "\n":
            if show_newline_glyph:
                el = ui.label('↵').classes(cls + ' newline-glyph')
                if on_click:
                    el.on('click', on_click)
            ui.html('').classes('flex-break')
        else:
            el = ui.label(tok).classes(cls)
            if on_click:
                el.on('click', on_click)

def get_ignore_pattern(field_name):
    
    if pattern := FIELD_MAP.get(field_name).get("ignore_pattern"):
        return pattern
    return None

def _render_diff_chunks(index:int, field_name: str, original: str, cand_text: str, size: str = 'normal'):
    ops = opcodes_for_diff(original, cand_text, get_ignore_pattern(field_name))
    with ui.row().classes('wrap items-start gap-1 ' + ('mono text-sm' if size == 'compact' else 'mono')):
        for idx, (tag, i1, i2, j1, j2) in enumerate(ops):
            if tag == 'equal':
                _emit_text_with_breaks(original[i1:i2], 'chunk equal', None, show_newline_glyph=False)
            elif tag == 'delete':
                _emit_text_with_breaks(original[i1:i2] or '', 'chunk delete',
                                    on_click=lambda e, idx=idx: apply_chunk(index, field_name, cand_text, idx, 'delete'),
                                    show_newline_glyph=True)
            elif tag == 'insert':
                _emit_text_with_breaks(cand_text[j1:j2] or '', 'chunk insert',
                                    on_click=lambda e, idx=idx: apply_chunk(index, field_name, cand_text, idx, 'insert'),
                                    show_newline_glyph=True)
            elif tag == 'replace':
                _emit_text_with_breaks(original[i1:i2] or '', 'chunk delete',
                                    on_click=lambda e, idx=idx: apply_chunk(index, field_name, cand_text, idx, 'delete'),
                                    show_newline_glyph=True)
                _emit_text_with_breaks(cand_text[j1:j2] or '', 'chunk insert',
                                    on_click=lambda e, idx=idx: apply_chunk(index, field_name, cand_text, idx, 'replace'),
                                    show_newline_glyph=True)
    return

def render_single_candidate(index:int, field_name: str, source: str, cand_text: str, container: ui.element):
    container.clear()
    original = store.records[index][field_name]
    with container:
        with ui.row().classes('items-center justify-between w-full'):
            ui.label(f'候选来源: {source}').classes('text-primary')
            with ui.row().classes('gap-2'):
                ui.button('应用全部变化', on_click=lambda: apply_all_from_candidate(index, field_name, cand_text)).props('dense flat')
                #ui.button('恢复为当前文本', on_click=lambda: refresh_record_view()).props('dense flat')
        _render_diff_chunks(index, field_name, original, cand_text, size='normal')


def render_diffs_for_field(index:int, field_name: str) -> None:
    # 确保存在容器
    if (index, field_name) not in field_diff_containers:
        field_diff_containers[(index,field_name)] = ui.column().classes('w-full')
    field_container = field_diff_containers[(index,field_name)]
    field_container.clear()
    rec_no = store.records[index][KEYNAME]
    candidates_dict = get_candidate(rec_no, field_name)
    if not candidates_dict:
        with field_container:
            ui.label('无候选数据').classes('muted')
        return
    with field_container:
        for source, cand_text in candidates_dict.items():
            with ui.card():
                inner = ui.column()
                render_single_candidate(index, field_name, source, cand_text, inner)

def render_compact_fields(container: ui.element):
    container.clear()
    compact_fields = [f for f, mode in FIELD_MODES.items() if mode == 'compact']
    if not compact_fields:
        return
    for offset in range(ITEMS_PER_PAGE):
        if current_index + offset >= get_visible_count():
            break
        real_idx = get_visible_index(current_index+offset)
        rec = store.records[real_idx]
        rec_no = rec[KEYNAME]
        sources: Dict[str, Dict[str, str]] = {}
        for f in compact_fields:
            for src, text in get_candidate(rec_no, f).items():
                sources.setdefault(src, {})[f] = text

        with container:
            with ui.row().classes('items-start gap-6'):
                for f in compact_fields:
                    inp = ui.input(value=rec.get(f, ''), placeholder=f'输入/编辑 {f} 文本…').classes('w-[280px] mono')
                    inp.on('blur', lambda e, field=f, idx=real_idx: update_field(idx, field, e.sender.value))
                    field_editors[(real_idx,f)] = inp

            for src, values in sources.items():
                with ui.row().classes('items-start gap-6'):
                    for f in compact_fields:
                        cand_text = values.get(f, '')
                        col = ui.column().classes('w-[280px]')
                        with col:
                            ui.label(f'{src}').classes('text-xs text-primary')
                            _render_diff_chunks(real_idx, f, rec.get(f, ''), cand_text, size='compact')
                            ui.button('应用全部', on_click=lambda f=f, t=cand_text, idx=real_idx: apply_all_from_candidate(idx, f, t)).props('dense flat')

# =============================
# Apply actions
# =============================

def apply_chunk(index:int, field_name: str, cand_text: str, opcode_index: int, action: str) -> None:
    original = store.records[index][field_name]
    new_value = apply_single_chunk(original, cand_text, opcode_index, action, get_ignore_pattern(field_name))
    store.set_field(index, field_name, new_value)
    field_editors[(index,field_name)].value = new_value
    if FIELD_MODES[field_name] == 'normal':
        render_diffs_for_field(index, field_name)
    else:
        render_compact_fields(compact_container)


def apply_all_from_candidate(index:int, field_name: str, cand_text: str) -> None:
    store.set_field(index, field_name, cand_text)
    field_editors[(index,field_name)].value = cand_text
    if FIELD_MODES[field_name] == 'normal':
        render_diffs_for_field(index, field_name)
    else:
        render_compact_fields(compact_container)

# =============================
# PDF Image Provider
# =============================

def merge_image(pdf, positions, direction="vertical", scale = False):
    """
    pdf: fitz.Document 对象
    positions: [{'page': 0, 'bbox': (x0, y0, x1, y1)}, ...]
    direction: "vertical" or "horizontal"
    """
    crops = []
    
    for pos in positions:
        page = pdf[pos["page"]-1]
        # 获取页面的所有图片
        img_list = page.get_images(full=True)
        if not img_list:
            continue
        
        # 取第一个图片
        xref = img_list[0][0]
        base_image = pdf.extract_image(xref)
        img_data = base_image["image"]
        img = Image.open(io.BytesIO(img_data))
        img = ImageOps.invert(img)
        
        # 注意：PDF 页面和图像坐标可能不一致，通常需要用 matrix 转换
        # 这里假设第一个图像覆盖整个页面（常见情况：扫描件）
        # 所以直接按照 bbox 比例在图像上裁剪
        page_rect = page.rect
        x0, y0, x1, y1 = pos["bbox"]
        
        # 转换 bbox 到图像坐标
        w, h = img.size
        if scale:
            crop_box = (
                int(x0 / page_rect.width * w),
                int(y0 / page_rect.height * h),
                int(x1 / page_rect.width * w),
                int(y1 / page_rect.height * h),
            )
        else:
            crop_box = (x0, y0, x1, y1)
        crop_img = img.crop(crop_box)
        crops.append(crop_img)

    # 拼接图片
    if not crops:
        return None
    
    if direction == "vertical":
        total_width = max(img.width for img in crops)
        total_height = sum(img.height for img in crops)
        result = Image.new("RGB", (total_width, total_height), (255, 255, 255))
        
        y_offset = 0
        for img in crops:
            result.paste(img, (0, y_offset))
            y_offset += img.height
    else:  # horizontal
        total_width = sum(img.width for img in crops)
        total_height = max(img.height for img in crops)
        result = Image.new("RGB", (total_width, total_height), (255, 255, 255))
        
        x_offset = 0
        for img in crops:
            result.paste(img, (x_offset, 0))
            x_offset += img.width
    
    return result

@app.get('/get_word_image')
def get_word_image(index: int, field: str = ''):
    if not image_source and not pdf_doc:
        return Response(status_code=404)
    if index < 0 or index >= len(store.records):
        return Response(status_code=404)
    rec = store.records[index]
    pdf_key = IMAGENAME
    if not field and not pdf_key:
        return Response(status_code=404)
    page_key = field + '_page' if field else 'page'
    bbox_key = field + '_bbox' if field else 'bbox'
    image_field = rec
    if pdf_key:
        position = rec.get(pdf_key)
    if isinstance(position,dict):
        if not page_key in rec or not bbox_key in rec:
            return Response(status_code=404)
        rect = fitz.Rect(*rec[bbox_key])
        page_no = rec[page_key] - 1
        pix = pdf_doc[page_no].get_pixmap(clip=rect, dpi=300)
        return Response(content=pix.tobytes('png'), media_type='image/png')
    elif isinstance(position,list):
        pix = merge_image(pdf_doc, position)
        buffer = io.BytesIO()
        pix.save(buffer, format="PNG")   # 保存到内存里
        png_bytes = buffer.getvalue()    # 取出 PNG 字节流

        return Response(content=png_bytes, media_type='image/png')
    else:
        return Response(status_code=404)
        
    

# =============================
# Layout with Config Loader
# =============================

def build_header():
    global main_column, record_status_label
    main_column = ui.column().classes('w-full')
    with ui.header().classes('items-center justify-between') as header:
        ui.label(TITLE).classes('text-h6')
        with ui.row().classes('items-center gap-2'):
            filter_options = [("无过滤", None)] + [
            (f_cfg["title"], make_filter(f_cfg)) for f_cfg in CONFIG.get("filters", [])
        ]
            def on_filter_change(e):
                idx = int(e.value)
                if idx == 0:
                    clear_filter()
                else:
                    f = filter_options[idx][1]
                    apply_filter(f)
            select_filter = ui.select({i: name for i, (name, _) in enumerate(filter_options)},
                                    label="过滤器", value=0,on_change=on_filter_change)
            
            ui.html('<input type="file" id="fileInput" style="display:none" onchange="window.loadConfigFromFile(this)" />')
            ui.button('⟵ 后退', on_click=lambda: set_current_index(current_index - ITEMS_PER_PAGE)).props('dense')
            ui.button('前进 ⟶', on_click=lambda: set_current_index(current_index + ITEMS_PER_PAGE)).props('dense')
            record_status_label = ui.label().classes('text-sm')
            if store.records:
                idx_input = ui.input(f'跳转到 {KEYNAME}', value=store.records[current_index][KEYNAME])#, min=store.records[0][KEYNAME], max=store.records[-1][KEYNAME])
                def goto():
                    no = idx_input.value
                    try:
                        idx = next(i for i, r in enumerate(store.records) if r.get(KEYNAME) == no)
                    except StopIteration:
                        ui.notify(f'找不到 {KEYNAME}={no}')
                        return
                    set_current_index(idx)
                ui.button('跳转', on_click=goto).props('dense')
            ui.separator().props('vertical')
            ui.button('保存', on_click=save_now).props('dense')


def build_layout():
    global compact_container, main_column, store
    

    with main_column:
        compact_container = ui.column().classes('w-full')
        for offset in range(ITEMS_PER_PAGE):
            if current_index + offset >= get_visible_count():
                break
            real_idx = get_visible_index(current_index+offset)
            rec = store.records[real_idx]
            ui.label(f"{KEYNAME}={rec.get(KEYNAME)}").classes('text-h6 text-primary')
            if IMAGENAME:
                ui.image(f"/get_word_image?index={real_idx}").style("max-width: 400px; height: auto;")
            for f in FIELDS:
                if FIELD_MODES[f] == 'normal':
                    if FIELD_PDF_KEYS.get(f) != None:
                        ui.image(f"/get_word_image?index={real_idx}&field={FIELD_PDF_KEYS[f]}")
                    ui.label(FIELD_LABELS[f]).classes('text-subtitle1')
                    ta = ui.textarea(value=rec.get(f, ''), placeholder=f'输入/编辑 {f} 文本…').classes('w-full mono')
                    ta.on('blur', lambda e, field_name=f, idx=real_idx: update_field(idx, field_name, e.sender.value))
                    field_editors[(real_idx,f)] = ta
                    field_diff_containers[(real_idx,f)] = ui.column().classes('w-full')
                    render_diffs_for_field(real_idx,f)
        render_compact_fields(compact_container)
        update_record_status()
# =============================
# Refresh
def refresh_record_view() -> None:
    main_column.clear()
    field_diff_containers.clear()
    field_editors.clear()
    with main_column:
        build_layout()


# =============================
# Init
# =============================

load_config(CONFIG_PATH)
build_header()
build_layout()
refresh_record_view()

ui.run(native=False, title=TITLE)